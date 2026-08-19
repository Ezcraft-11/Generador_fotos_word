
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def insertar_imagenes_2x2_horizontal(carpeta_origen, archivo_salida="Imagenes_Organizadas_2x2.docx"):
    # 1. Crear un documento nuevo
    doc = Document()
    
    # 2. Configurar la orientación de la página a HORIZONTAL (Landscape)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    
    # En python-docx es necesario intercambiar ancho y alto al cambiar la orientación
    nuevo_ancho, nuevo_alto = section.page_height, section.page_width
    section.page_width = nuevo_ancho
    section.page_height = nuevo_alto

    # 3. Configurar márgenes estrechos (0.5 pulgadas) para maximizar el espacio
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Extensiones de imagen permitidas
    extensiones_validas = ('.png', '.jpg', '.jpeg', '.webp', '.bmp')
    
    if not os.path.exists(carpeta_origen):
        print(f"Error: La carpeta '{carpeta_origen}' no existe.")
        return

    # Filtrar y ordenar las imágenes de la carpeta
    archivos = [f for f in os.listdir(carpeta_origen) if f.lower().endswith(extensiones_validas)]
    archivos.sort()

    if not archivos:
        print("No se encontraron imágenes en la carpeta especificada.")
        return

    # 4. Procesar las fotos en bloques de 4
    for i in range(0, len(archivos), 4):
        grupo = archivos[i:i+4]
        
        # Crear tabla de 2 filas x 2 columnas para el formato 2x2
        tabla = doc.add_table(rows=2, cols=2)
        tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabla.autofit = False

        # Coordenadas dentro de la tabla 2x2
        posiciones = [(0, 0), (0, 1), (1, 0), (1, 1)]
        
        for idx, nombre_foto in enumerate(grupo):
            row_idx, col_idx = posiciones[idx]
            celda = tabla.cell(row_idx, col_idx)
            
            # Ajustar párrafos dentro de la celda
            p = celda.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            ruta_foto = os.path.join(carpeta_origen, nombre_foto)
            
            # Ancho estándar de 4.5 pulgadas para que 2 fotos quepan perfectamente a lo ancho
            run = p.add_run()
            run.add_picture(ruta_foto, width=Inches(4.5))

        # 5. Insertar salto de página si quedan más fotos pendientes
        if i + 4 < len(archivos):
            doc.add_page_break()

    # Guardar documento
    doc.save(archivo_salida)
    print(f"¡Proceso completado! Archivo guardado como '{archivo_salida}'.")

# --- USO DEL SCRIPT ---
if __name__ == "__main__":
    # Sustituye './mis_fotos' por la ruta de tu carpeta
    carpeta_fotos = "./mis_fotos"
    insertar_imagenes_2x2_horizontal(carpeta_fotos)
