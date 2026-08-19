
import os
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import time
def generar_reporte_limpio(carpeta_fotos):
    if not os.path.exists(carpeta_fotos):
        print(f"Error: No encuentro la carpeta '{carpeta_fotos}'")
        return

    # Buscar fotos
    fotos = [f for f in os.listdir(carpeta_fotos) if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    fotos.sort()

    if not fotos:
        print("No hay fotos.")
        return

    # Agrupar de 4 en 4
    grupos = [fotos[i:i + 4] for i in range(0, len(fotos), 4)]

    for idx, grupo in enumerate(grupos):
        doc = Document()
        
        # --- PONER LA HOJA HORIZONTAL ---
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # Márgenes muy pequeños para que luzcan las fotos
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.1)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        # Tabla de 2x2 sin bordes
        tabla = doc.add_table(rows=2, cols=2)
        
        for i, nombre_foto in enumerate(grupo):
            fila = i // 2
            columna = i % 2
            celda = tabla.rows[fila].cells[columna]
            
            ruta_foto = os.path.join(carpeta_fotos, nombre_foto)
            parrafo = celda.paragraphs[0]
            run = parrafo.add_run()
            
            # Ajuste de tamaño para horizontal (4.5 pulgadas es ideal)
            run.add_picture(ruta_foto, width=Inches(4.5))

        # Guardar
        nombre_final = f"reporte_horizontal_{idx + 1}.docx"
        doc.save(nombre_final)
        print(f"✅ ¡Listo! Se creó {nombre_final} sin letras y acostado.")
        ruta_completa = os.path.abspath(nombre_final)

        # try:
            # print(f"enviar a la impresora")
            # os.startfile(ruta_completa, "print")
            # time.sleep(5)
        # except Exception as e:
            # print(f"no se pudo imprimir")

# Ejecutar usando tu carpeta
generar_reporte_limpio("./Fotos_JPG")

